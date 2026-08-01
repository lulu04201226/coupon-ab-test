from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT = Path(__file__).resolve().parents[1]
TABLE_DIR = PROJECT / "outputs" / "tables"
CHART_DIR = PROJECT / "outputs" / "charts"
REPORT_DIR = PROJECT / "report"
REPORT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT = REPORT_DIR / "网约车优惠券AB测试分析报告.docx"

BLUE = "2563EB"
DARK_BLUE = "1E3A5F"
NAVY = "172B4D"
GRAY = "667085"
LIGHT_BLUE = "EAF2FF"
LIGHT_GRAY = "F2F4F7"
GREEN = "15803D"
RED = "B42318"
GOLD = "9A6700"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_font(run, size=10.5, bold=False, color="222222", italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)


def set_cell_fill(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(round(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def style_table(table, widths, header_fill=LIGHT_GRAY):
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, widths)
    set_repeat_header(table.rows[0])
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_fill(cell, header_fill)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    set_font(run, size=9, bold=(r_idx == 0), color=NAVY if r_idx == 0 else "222222")


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, bold=True, color=NAVY)
        r = p.add_run(text[len(bold_prefix):])
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullets(doc, items, style="List Bullet"):
    for item in items:
        p = doc.add_paragraph(style=style)
        r = p.add_run(item)
        set_font(r)


def add_small_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_font(r, size=8.5, color=GRAY, italic=True)
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE, label_color=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}｜")
    set_font(r, bold=True, color=label_color)
    r = p.add_run(text)
    set_font(r)
    return table


def add_picture(doc, filename, caption):
    path = CHART_DIR / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(6.15))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_font(r, size=8.5, color=GRAY)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_font(run, size=8.5, color=GRAY)


summary = pd.read_csv(TABLE_DIR / "group_summary.csv", index_col=0)
tests = pd.read_csv(TABLE_DIR / "paired_test_results.csv").set_index("metric")
method = pd.read_csv(TABLE_DIR / "gmv_method_comparison.csv")
power = pd.read_csv(TABLE_DIR / "power_analysis.csv").iloc[0]
economics = pd.read_csv(TABLE_DIR / "economics_summary.csv").iloc[0]
robustness = pd.read_csv(TABLE_DIR / "gmv_robustness_checks.csv")
coupon_trend = pd.read_csv(TABLE_DIR / "coupon_trend.csv")

control = summary.loc["control"]
experiment = summary.loc["experiment"]
gmv = tests.loc["gmv"]
coupon_cost = tests.loc["coupon_cost"]
completion = tests.loc["completion_rate"]
cancel = tests.loc["cancel_rate"]
aov = tests.loc["aov"]
roi = tests.loc["roi"]

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.85)
sec.bottom_margin = Inches(0.85)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.header_distance = Inches(0.42)
sec.footer_distance = Inches(0.42)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, DARK_BLUE, 11, 5),
    ("Heading 3", 11.5, DARK_BLUE, 8, 4),
]:
    st = doc.styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = rgb(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for name in ["List Bullet", "List Number", "List Bullet 2"]:
    st = doc.styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.12

header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = header.add_run("网约车优惠券 A/B 测试分析｜项目报告")
set_font(r, size=8.5, color=GRAY)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("数据范围：2019年1月1日—29日  ·  第 ")
set_font(r, size=8.5, color=GRAY)
add_page_field(footer)
r = footer.add_run(" 页")
set_font(r, size=8.5, color=GRAY)

# Cover
doc.add_paragraph().paragraph_format.space_after = Pt(68)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("A/B TEST ANALYSIS")
set_font(r, size=11, bold=True, color=BLUE)
p.paragraph_format.space_after = Pt(18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("网约车优惠券 A/B 测试分析")
set_font(r, size=29, bold=True, color=NAVY)
p.paragraph_format.space_after = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("提高优惠券力度是否值得推广？")
set_font(r, size=15, color=GRAY)
p.paragraph_format.space_after = Pt(48)

table = doc.add_table(rows=4, cols=2)
cover_rows = [
    ("分析对象", "对照组与实验组按日期配对的29天汇总数据"),
    ("主要指标", "GMV、订单量、优惠券成本、ROI、完成率、取消率"),
    ("核心方法", "配对t检验、Wilcoxon、Bootstrap、功效与MDE分析"),
    ("数据来源", "test.xlsx"),
]
for i, (label, value) in enumerate(cover_rows):
    table.cell(i, 0).text = label
    table.cell(i, 1).text = value
style_table(table, [1.35, 5.15], header_fill=LIGHT_BLUE)
for row in table.rows:
    set_cell_fill(row.cells[0], LIGHT_BLUE)
    row.cells[0].paragraphs[0].runs[0].bold = True
doc.add_paragraph().paragraph_format.space_after = Pt(24)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("项目报告 · 咨询式精简结构")
set_font(r, size=9, color=GRAY)

doc.add_page_break()

# Executive summary first
add_heading(doc, "Executive Summary｜结论速览", 1)
add_callout(
    doc,
    "决策建议",
    "不建议直接推广当前优惠券路径。实验组提高了补贴投入，但GMV和客单价显著下降；建议停止直接放量，控制券额衰减路径后进行第二轮受控实验。",
    fill="FFF4E5",
    label_color=GOLD,
)
add_body(
    doc,
    f"收益指标转弱。实验组日均GMV比对照组低 {abs(gmv['mean_diff']):,.0f}（{gmv['relative_diff']:.2%}），"
    f"日均客单价下降 {abs(aov['mean_diff']):.2f}（{aov['relative_diff']:.2%}）。",
    bold_prefix="收益指标转弱。",
)
add_body(
    doc,
    f"结论依赖正确的配对结构。忽略日期配对的Welch检验p={method.iloc[0]['p_value']:.4f}，"
    f"会得到“没有差异”；按日期配对后p={gmv['paired_t_p']:.4f}，95%置信区间为"
    f"[{gmv['ci_low']:,.0f}, {gmv['ci_high']:,.0f}]，且方向为下降。",
    bold_prefix="结论依赖正确的配对结构。",
)
add_body(
    doc,
    f"补贴经济性不成立。实验组日均优惠券成本增加 {coupon_cost['relative_diff']:.2%}，"
    f"绝对ROI下降 {abs(roi['relative_diff']):.2%}。实验期总GMV减少 {abs(economics['delta_gmv']):,.0f}，"
    f"同时优惠券成本增加 {economics['delta_coupon_cost']:,.0f}。",
    bold_prefix="补贴经济性不成立。",
)
add_body(
    doc,
    f"体验护栏小幅改善。日均完成率提高 {completion['mean_diff']:.2%}，"
    f"取消率降低 {abs(cancel['mean_diff']):.2%}；改善幅度不足以抵消GMV和客单价下降。",
    bold_prefix="体验护栏小幅改善。",
)
add_body(
    doc,
    f"检验敏感性充足但效应存在时间异质性。当前29对样本下，80%功效对应的GMV MDE约"
    f"{power['mde_gmv_relative']:.2%}；负向效果主要集中在后半程，因此需在固定处理强度下复验。",
    bold_prefix="检验敏感性充足但效应存在时间异质性。",
)

add_heading(doc, "核心指标一览", 2)
table = doc.add_table(rows=1, cols=5)
for i, h in enumerate(["指标", "对照组日均", "实验组日均", "相对变化", "配对检验"]):
    table.cell(0, i).text = h
for metric, formatter in [
    ("gmv", lambda x: f"{x:,.0f}"),
    ("coupon_cost", lambda x: f"{x:,.1f}"),
    ("completion_rate", lambda x: f"{x:.2%}"),
    ("cancel_rate", lambda x: f"{x:.2%}"),
    ("aov", lambda x: f"{x:,.2f}"),
]:
    row = tests.loc[metric]
    cells = table.add_row().cells
    values = [
        row["metric_cn"],
        formatter(row["control_mean"]),
        formatter(row["experiment_mean"]),
        f"{row['relative_diff']:+.2%}",
        f"p={row['paired_t_p']:.4f}",
    ]
    for i, value in enumerate(values):
        cells[i].text = value
style_table(table, [1.25, 1.35, 1.35, 1.1, 1.45])

# 01
add_heading(doc, "01｜项目背景与业务问题", 1)
add_heading(doc, "1.1 提高优惠券力度，必须同时验证增长与效率", 2)
add_body(
    doc,
    "网约车平台希望通过提高优惠券力度刺激用户请求与成交，但补贴只有在带来足够新增GMV、订单或用户价值时才值得推广。"
    "本项目将增长指标、成本效率和体验护栏放在同一决策框架中。"
)
add_heading(doc, "1.2 本报告回答五个问题", 2)
add_bullets(
    doc,
    [
        "两组是否可比，应该使用独立样本还是日期级配对分析？",
        "提高优惠券力度是否拉动请求量、完成订单数与GMV？",
        "新增收益能否覆盖新增补贴，ROI是否改善？",
        "策略是否以取消率上升或完成率下降为代价？",
        "当前实验的检验敏感性是否足够支撑结论？",
    ],
)

# 02
add_heading(doc, "02｜数据说明与实验方案设计", 1)
add_small_note(
    doc,
    "说明：原始数据仅包含实验结果，不含方案文档、分流规则及上线记录。本节依据现有字段和标准A/B测试流程补充完整项目链路；实验日期来自数据，其余方案设计为建议，不代表实际执行记录。",
)
add_heading(doc, "2.1 数据事实：29天×2组形成完整日期配对", 2)
table = doc.add_table(rows=1, cols=3)
for i, h in enumerate(["项目", "结果", "判断"]):
    table.cell(0, i).text = h
for row in [
    ("数据周期", "2019-01-01—2019-01-29", "覆盖29个连续日期"),
    ("样本结构", "对照组29条、实验组29条", "同日期一一配对"),
    ("数据质量", "0缺失、0重复", "可直接进入分析"),
    ("分析粒度", "日期×实验组别", "不能替代用户级日志"),
]:
    cells = table.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value
style_table(table, [1.25, 2.2, 3.05])

add_heading(doc, "2.2 实验目标、指标与假设（方案补充）", 2)
add_bullets(
    doc,
    [
        "实验目标：判断提高优惠券力度能否增加GMV和订单，并验证新增收益能否覆盖新增补贴。",
        "主要指标：GMV；经济性指标：绝对ROI与增量价值。",
        "诊断指标：requests、trips、客单价。",
        "护栏指标：完成率、取消率。",
        "H₀：实验组与对照组的日期级核心指标均值差为0；H₁：均值差不为0。",
    ],
)
add_heading(doc, "2.3 推荐分流与灰度路径（方案补充）", 2)
add_bullets(
    doc,
    [
        "采用用户ID稳定哈希分流，实验期间保持用户组别不变。",
        "5%灰度：验证分流、领券、核销、指标监控和回滚链路。",
        "20%放量：重点观察取消率、完成率和补贴成本。",
        "50%正式实验：达到事前样本量与最短周期后进行正式判断。",
        "若取消率上升、完成率下降或成本超过预算阈值，则触发回滚。",
    ],
)
add_heading(doc, "2.4 周期与样本量设计（方案补充）", 2)
add_body(
    doc,
    f"正式实验建议至少覆盖两个完整业务周期，并预先锁定显著性水平5%、目标功效80%和主要指标。"
    f"以当前数据的日期配对波动估算，29对样本下GMV的MDE约为日均{power['mde_gmv_relative']:.2%}；"
    f"若真实效应等于本次观测效应，约{int(power['pairs_needed_80_power'])}对日期可达到80%功效。"
)
add_small_note(
    doc,
    f"基于已观测效应量计算的事后功效约{power['posthoc_power']:.1%}，仅用于敏感性说明，不作为实验成功的独立证据。正式上线应使用历史基线方差进行事前测算。",
)

# 03
add_heading(doc, "03｜基线结构决定必须采用日期配对分析", 1)
add_body(
    doc,
    "两组在同一日期的请求量、订单量和GMV同步涨跌，日期间波动远大于组间差异。"
    "因此分析重点不是跨日期比较两个总体，而是在同一天内比较实验组和对照组。"
)
add_picture(doc, "01_daily_trends.png", "图1｜两组核心指标每日趋势：同步波动支持日期级配对分析。")
add_heading(doc, "3.1 处理强度并非固定", 2)
add_body(
    doc,
    "两组每单优惠券金额都随时间明显下降。对照组拟合日斜率为"
    f"{coupon_trend.loc[coupon_trend['group']=='control','daily_slope'].iloc[0]:.4f}，实验组为"
    f"{coupon_trend.loc[coupon_trend['group']=='experiment','daily_slope'].iloc[0]:.4f}。"
    "这意味着实验不是固定券额差异的简单干预，后续结论应解释为“当前优惠券路径”的总体效果。"
)
add_picture(doc, "06_coupon_intensity_trend.png", "图2｜每单优惠券金额趋势：两组券额均随实验推进下降。")

# 04
add_heading(doc, "04｜实验效果：体验改善，但增长与效率转弱", 1)
add_heading(doc, "4.1 请求和订单没有形成显著增量", 2)
requests = tests.loc["requests"]
trips = tests.loc["trips"]
add_body(
    doc,
    f"实验组日均请求量下降{abs(requests['relative_diff']):.2%}（p={requests['paired_t_p']:.3f}），"
    f"日均完成订单量下降{abs(trips['relative_diff']):.2%}（p={trips['paired_t_p']:.3f}），均未达到5%显著性水平。"
)
add_heading(doc, "4.2 GMV下降来自客单价，而非订单量大幅减少", 2)
add_body(
    doc,
    f"实验组日均GMV下降{abs(gmv['relative_diff']):.2%}，客单价下降{abs(aov['relative_diff']):.2%}。"
    "完成订单量基本持平，因此GMV损失主要来自每笔完成订单贡献下降，而不是订单规模明显萎缩。"
)
add_picture(doc, "03_metric_lifts.png", "图3｜实验组相对对照组的日均指标变化：增长与效率指标走弱，体验护栏改善。")
add_heading(doc, "4.3 补贴成本增加，增量经济性为负", 2)
add_body(
    doc,
    f"实验组日均优惠券成本增加{coupon_cost['relative_diff']:.2%}，绝对ROI下降{abs(roi['relative_diff']):.2%}。"
    f"整个实验周期实验组GMV比对照组少{abs(economics['delta_gmv']):,.0f}，优惠券成本却多{economics['delta_coupon_cost']:,.0f}。"
)
add_callout(
    doc,
    "解读边界",
    "由于增量优惠券成本分母很小且增量GMV为负，不应把增量ROI比值外推为稳定参数。更可靠的业务判断是：当前策略增加了补贴，但没有形成正向GMV增量。",
    fill="FFF4E5",
    label_color=GOLD,
)
add_heading(doc, "4.4 完成率和取消率改善，但不足以支持推广", 2)
add_body(
    doc,
    f"实验组日均完成率提高{completion['mean_diff']:.2%}，取消率降低{abs(cancel['mean_diff']):.2%}，"
    "均通过配对检验和Holm校正。体验端改善是真实信号，但商业目标GMV和ROI同时恶化，不能仅凭护栏改善推广。"
)

# 05
add_heading(doc, "05｜配对检验揭示被日期波动掩盖的负向GMV效果", 1)
add_heading(doc, "5.1 错误忽略配对会得到“没有差异”", 2)
add_body(
    doc,
    f"将两组29天数据误作独立样本时，Welch检验p={method.iloc[0]['p_value']:.4f}；"
    f"按日期配对后p={method.iloc[1]['p_value']:.4f}。配对分析消除了两组共有的日期波动，"
    "使组间系统性差异从噪声中显现。"
)
add_picture(doc, "04_method_reversal.png", "图4｜GMV检验方法对显著性判断的影响：配对结构改变结论。")
add_heading(doc, "5.2 GMV差值显著为负，非参数检验结论一致", 2)
add_body(
    doc,
    f"实验组－对照组的GMV日均差为{gmv['mean_diff']:,.0f}，95%置信区间为"
    f"[{gmv['ci_low']:,.0f}, {gmv['ci_high']:,.0f}]；Cohen's dz={gmv['cohen_dz']:.2f}。"
    f"Wilcoxon检验p={gmv['wilcoxon_p']:.4f}，与配对t检验方向一致。"
)
add_picture(doc, "02_gmv_paired_difference.png", "图5｜每日GMV配对差值：多数后半程日期表现为负。")
add_heading(doc, "5.3 稳健性检查支持负向结论，但提示时间异质性", 2)
boot = robustness.loc[robustness["check"] == "bootstrap_20000"].iloc[0]
early = robustness.loc[robustness["check"] == "first_half"].iloc[0]
late = robustness.loc[robustness["check"] == "second_half"].iloc[0]
add_body(
    doc,
    f"20,000次Bootstrap得到GMV均值差95%区间[{boot['ci_low']:,.0f}, {boot['ci_high']:,.0f}]，仍完全低于0。"
    f"但前14天均值差为{early['mean_diff']:,.0f}（p={early['p_value']:.3f}），后15天为"
    f"{late['mean_diff']:,.0f}（p<0.001）。因此负向效果并非全周期均匀，可能与券额路径、用户构成或其他时间因素有关。"
)
add_heading(doc, "5.4 当前样本能识别约0.86%的日均GMV变化", 2)
add_body(
    doc,
    f"在29对日期、双侧α=5%、目标功效80%的条件下，GMV最小可检测效应约为日均"
    f"{power['mde_gmv_relative']:.2%}（约{power['mde_gmv_absolute']:,.0f}）。"
    "本次观测到的1.26%下降超过该阈值，说明GMV结论不是简单的“样本不足”。"
)
add_picture(doc, "05_power_and_mde.png", "图6｜功效与MDE敏感性：样本增加会提高功效并降低可检测效应门槛。")
add_heading(doc, "5.5 统计显著不等于值得推广", 2)
add_body(
    doc,
    "GMV下降不仅统计显著，也具有明确的负向业务含义；完成率和取消率改善虽显著，但绝对幅度较小。"
    "推广判断必须以主要收益指标和补贴经济性为核心，再结合护栏指标确认风险。"
)

# 06
add_heading(doc, "06｜结论与行动建议", 1)
add_heading(doc, "6.1 核心结论", 2)
add_bullets(
    doc,
    [
        "当前优惠券路径没有拉动请求或订单，反而使日均GMV下降1.26%。",
        "GMV下降主要由客单价降低驱动，且优惠券成本增加、绝对ROI下降。",
        "完成率提高、取消率下降，说明体验侧可能存在正向作用，但不足以抵消收益损失。",
        "日期配对是正确分析结构；忽略配对会把显著负向结果误判为“无差异”。",
        "效果在前后半程明显不同，处理强度随时间变化，需通过二次受控实验验证机制。",
    ],
)
add_heading(doc, "6.2 P0｜停止直接放量，重新定义推广门槛", 2)
add_bullets(
    doc,
    [
        "不按当前券额路径全量推广。",
        "以GMV或净增量价值作为主要成功指标，而不是仅观察订单完成率。",
        "推广门槛同时要求：主要指标达到预设最小业务提升、增量价值为正、护栏指标不恶化。",
    ],
)
add_heading(doc, "6.3 P1｜开展固定处理强度的第二轮实验", 2)
add_bullets(
    doc,
    [
        "固定实验期内的券额差，或按券额档位预先分层，避免处理强度随时间漂移。",
        "保留用户ID稳定分流，并记录实验曝光、领券、核销和订单链路。",
        "上线前完成SRM检查、事前样本量设计和停止规则。",
        "增加券使用率、用户分层、毛利率与长期复购指标，解释客单价下降机制。",
        "预先注册前后半程或业务周期异质性分析，避免事后切片。",
    ],
)
add_heading(doc, "6.4 后续仍需回答的问题", 2)
add_bullets(
    doc,
    [
        "实验组客单价下降，是低价订单被激活，还是高价值订单被补贴替代？",
        "后半程GMV下降是否与券额差、用户构成或外部时间因素有关？",
        "完成率与取消率改善能否带来长期留存价值？",
        "纳入毛利、平台抽佣及运营成本后，策略的利润增量是多少？",
    ],
)

# 07
add_heading(doc, "07｜局限性与数据口径", 1)
add_bullets(
    doc,
    [
        "数据为日期级汇总结果，并非用户级实验日志，无法核验随机化质量、SRM和重复曝光。",
        "每单优惠券金额随时间变化，组间处理强度并非恒定，效应应解释为当前策略路径的总体结果。",
        "前后半程差异属于事后发现，需通过预先设计的二次实验验证。",
        "数据仅覆盖29天，不能直接外推长期留存和长期利润效果。",
        "ROI只使用GMV与优惠券成本，不包含抽佣、司机激励和其他运营成本，不等同于利润ROI。",
        "同时检验多个指标可能增加假阳性风险；本文补充Holm校正，并以GMV为主要决策指标。",
    ],
)

# Appendix
add_heading(doc, "附录A｜数据字典", 1)
table = doc.add_table(rows=1, cols=3)
for i, h in enumerate(["字段", "类型", "业务含义"]):
    table.cell(0, i).text = h
for row in [
    ("date", "日期", "观测日期"),
    ("group", "文本", "control / experiment"),
    ("requests", "整数", "订单请求数"),
    ("gmv", "数值", "成交总额"),
    ("coupon per trip", "数值", "每单优惠券金额"),
    ("trips", "整数", "完成订单数"),
    ("canceled requests", "整数", "取消请求数"),
]:
    cells = table.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value
style_table(table, [1.55, 1.05, 3.9])

add_heading(doc, "附录B｜指标与方法口径", 1)
add_bullets(
    doc,
    [
        "优惠券成本 = coupon per trip × trips。",
        "绝对ROI = GMV ÷ 优惠券成本。",
        "完成率 = trips ÷ requests。",
        "取消率 = canceled requests ÷ requests。",
        "客单价 = GMV ÷ trips。",
        "配对差值 = 同日期实验组指标 − 对照组指标。",
        "Cohen's dz = 配对差值均值 ÷ 配对差值标准差。",
        "MDE：在给定α、功效、样本量和配对差值波动下可识别的最小效应。",
    ],
)
add_small_note(
    doc,
    "分析代码、完整统计结果表及图表均保存在本项目目录中。所有报告数字均由 test.xlsx 重新计算生成。",
)

doc.core_properties.title = "网约车优惠券A/B测试分析报告"
doc.core_properties.subject = "提高优惠券力度是否值得推广"
doc.core_properties.author = ""
doc.save(OUTPUT)
print(OUTPUT)
